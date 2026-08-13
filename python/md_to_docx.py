# -*- coding: utf-8 -*-
"""
md_to_docx.py
Convert Dissertation_Draft_v2.md to a formatted Word document.

Author : Ghid Albazrkan
Degree : MSc GIS, University of Aberdeen (GG5910/GG5912)
"""

import os
import re
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE      = r"C:\Users\HP\Desktop\proposal& dissertation"
MD_FILE   = os.path.join(BASE, "dissertation", "Dissertation_Draft_v2.md")
OUT_FILE  = os.path.join(BASE, "dissertation", "Dissertation_Draft_v2.docx")

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins: 3cm top/bottom, 2.5cm left/right
section = doc.sections[0]
section.top_margin    = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(18)  # 1.5 line spacing

# ── Style helpers ──────────────────────────────────────────────────────────────
def set_heading_style(para, level):
    sizes = {1: 16, 2: 14, 3: 13, 4: 12}
    bold  = {1: True, 2: True, 3: True, 4: True}
    para.style = doc.styles[f"Heading {level}"]
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(sizes.get(level, 12))
        run.font.bold = bold.get(level, True)
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_page_number():
    """Add page number at bottom centre."""
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    run2 = para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    run2._r.append(instrText)
    run3 = para.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)

add_page_number()

# ── Inline markdown parser ────────────────────────────────────────────────────
def add_inline(para, text):
    """Add text to paragraph handling **bold**, *italic*, `code`."""
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            if part:
                para.add_run(part)

# ── Table parser ──────────────────────────────────────────────────────────────
def parse_table(lines, start_idx):
    """Parse markdown table lines, return (table_lines, next_idx)."""
    table_lines = []
    idx = start_idx
    while idx < len(lines):
        line = lines[idx].strip()
        if line.startswith("|"):
            table_lines.append(line)
            idx += 1
        else:
            break
    return table_lines, idx

def add_table(table_lines):
    """Add a Word table from markdown table lines."""
    # Filter out separator rows (---|---|---)
    data_rows = [l for l in table_lines if not re.match(r"^\|[-| :]+\|$", l.strip())]
    if not data_rows:
        return
    rows = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < ncols:
                cell = tbl.rows[r_idx].cells[c_idx]
                cell.text = ""
                para = cell.paragraphs[0]
                add_inline(para, cell_text)
                if r_idx == 0:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()  # spacing after table

# ── Main parser ───────────────────────────────────────────────────────────────
print(f"Reading: {MD_FILE}")
with open(MD_FILE, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    raw = lines[i]
    line = raw.rstrip("\n")

    # Skip YAML front matter
    if line.strip() == "---" and i == 0:
        i += 1
        while i < len(lines) and lines[i].strip() != "---":
            i += 1
        i += 1
        continue

    # Horizontal rule
    if re.match(r"^---+$", line.strip()) or re.match(r"^\*\*\*+$", line.strip()):
        doc.add_paragraph("_" * 60)
        i += 1
        continue

    # Headings
    m = re.match(r"^(#{1,4})\s+(.*)", line)
    if m:
        level = len(m.group(1))
        heading_text = m.group(2).strip()
        para = doc.add_heading(level=min(level, 4))
        para.clear()
        add_inline(para, heading_text)
        set_heading_style(para, min(level, 4))
        i += 1
        continue

    # Table
    if line.strip().startswith("|"):
        table_lines, i = parse_table(lines, i)
        add_table(table_lines)
        continue

    # Bullet list
    m = re.match(r"^(\s*[-*])\s+(.*)", line)
    if m:
        para = doc.add_paragraph(style="List Bullet")
        add_inline(para, m.group(2).strip())
        para.paragraph_format.line_spacing = Pt(18)
        i += 1
        continue

    # Numbered list
    m = re.match(r"^\s*\d+\.\s+(.*)", line)
    if m:
        para = doc.add_paragraph(style="List Number")
        add_inline(para, m.group(1).strip())
        para.paragraph_format.line_spacing = Pt(18)
        i += 1
        continue

    # Empty line
    if line.strip() == "":
        i += 1
        continue

    # Normal paragraph
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = Pt(18)
    add_inline(para, line.strip())
    i += 1

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"\nSaved: {OUT_FILE}")
print("Done.")
